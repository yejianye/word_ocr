import re
import os
import base64
import hashlib
import io
from pprint import pprint
from textwrap import dedent


import cv2
import boto3
import numpy as np

from docx import Document
from docx.shared import RGBColor, Pt
from openai import OpenAI
from joblib import Memory

import imgutil
from util import cache, generate_file_hash, strip_format_quote
from llm import llm_completion, llm_image_completion

# ==== COMMON ====
def preprocess_image(image_path_or_file):
    image_hash = generate_file_hash(image_path_or_file)
    processed_image = f"/tmp/img_{image_hash}.jpg"
    imgutil.preprocess_image(image_path_or_file, processed_image)
    return processed_image

# ==== OCR with Textract ====
def find_highlighted_regions(image_path_or_image, region_min_size=200, output_path=None):
    if isinstance(image_path_or_image, str):
        image = cv2.imread(image_path_or_image)
    else:
        image = image_path_or_image

    # Convert the image to HSV color space for better color segmentation
    hsv_image = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)

    # Define the range for orange highlights in HSV
    orange_lower = np.array([10, 80, 100], dtype=np.uint8)
    orange_upper = np.array([60, 255, 255], dtype=np.uint8)

    # Create masks for orange and yellow regions
    orange_mask = cv2.inRange(hsv_image, orange_lower, orange_upper)

    # Find contours of the highlighted regions
    contours, _ = cv2.findContours(orange_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # Draw bounding boxes around the highlighted regions
    regions = []
    for contour in contours:
        # Get the bounding box for each contour
        x, y, w, h = cv2.boundingRect(contour)
        # Draw a rectangle around the highlighted area
        if w * h > region_min_size:
            regions.append((x, y, w, h))
            if output_path:
                cv2.rectangle(image, (x, y), (x + w, y + h), (0, 255, 0), 2)

    if output_path:
        cv2.imshow('Highlighted Regions', image)
        cv2.waitKey(0)
        cv2.destroyAllWindows()
        cv2.imwrite(output_path, image)

    return regions

def intersect_area(box1, box2):
    x1, y1, w1, h1 = box1
    x2, y2, w2, h2 = box2
    x_overlap = max(0, min(x1 + w1, x2 + w2) - max(x1, x2))
    y_overlap = max(0, min(y1 + h1, y2 + h2) - max(y1, y2))
    return x_overlap * y_overlap

def sanitize_text(text):
    # Use regex to extract only alphabets and hyphens from the text
    match = re.search(r'[a-zA-Z- ]+', text)
    if match:
        return match.group().strip()
    return None


@cache
def get_text_with_textract(image_path_or_image):
    if isinstance(image_path_or_image, str):
        image = cv2.imread(image_path_or_image)
    else:
        image = image_path_or_image

    height, width = image.shape[:2]
    _, image_data = cv2.imencode(".jpg", image)
    image_bytes = image_data.tobytes()

    textract = boto3.client('textract')
    response = textract.detect_document_text(
        Document={'Bytes': image_bytes}
    )
    # Process each detected block
    words = []
    full_text = ""
    idx = 0
    for block in response['Blocks']:
        if block['BlockType'] == 'LINE':
            full_text += block['Text'] + "\n"
        if block['BlockType'] == 'WORD':  # Get bounding boxes for words only
            text = block['Text']
            bounding_box = block['Geometry']['BoundingBox']
            bounding_box = (
                int(bounding_box['Left'] * width),
                int(bounding_box['Top'] * height),
                int(bounding_box['Width'] * width),
                int(bounding_box['Height'] * height)
            )
            confidence = block['Confidence'] / 100
            words.append({
                'text': text,
                'idx': idx,
                'confidence': confidence,
                'bounding_box': bounding_box
            })
            idx += 1
    return {'full_text': full_text, 'words': words}

def merge_bounding_boxes(boxes):
    if not boxes:
        return None
    if len(boxes) == 1:
        return boxes[0]
        
    # Initialize with first box
    x1, y1, w1, h1 = boxes[0]
    min_x = x1
    min_y = y1
    max_x = x1 + w1
    max_y = y1 + h1
    
    # Find min/max coordinates across all boxes
    for box in boxes[1:]:
        x, y, w, h = box
        min_x = min(min_x, x)
        min_y = min(min_y, y)
        max_x = max(max_x, x + w)
        max_y = max(max_y, y + h)
    
    # Calculate width and height of merged box
    width = max_x - min_x
    height = max_y - min_y
    
    return (min_x, min_y, width, height)

def calculate_text_bounding_box(words, margin=20):
    """
    Calculate the merged bounding box for all words with a safety margin.
    
    Args:
        words (list): List of word dictionaries containing bounding box information
        margin (int, optional): Safety margin to add around the bounding box. Defaults to 20.
    
    Returns:
        tuple: (x1, y1, x2, y2) coordinates of the expanded bounding box
    """
    text_bounding_box = merge_bounding_boxes([word['bounding_box'] for word in words])
    return (
        text_bounding_box[0] - margin,
        text_bounding_box[1] - margin,
        text_bounding_box[2] + margin,
        text_bounding_box[3] + margin
    )

def filter_regions_outside_bounding_box(regions, bounding_box):
    filtered_regions = []
    for region in regions:
        x, y, w, h = region
        if x >= bounding_box[0] and y >= bounding_box[1] and x + w <= bounding_box[0] + bounding_box[2] and y + h <= bounding_box[1] + bounding_box[3]:
            filtered_regions.append(region)
    return filtered_regions

def extract_highlighted_words_from_image(image_path_or_file, area_threshold=200):
    processed_image = preprocess_image(image_path_or_file)
    # Load the image using OpenCV
    debug_image = cv2.imread(processed_image)
    hl_regions = find_highlighted_regions(processed_image)
    # Use pytesseract to get the bounding box information
    result = get_text_with_textract(processed_image)
    text_bounding_box = calculate_text_bounding_box(result['words'])
    # Filter out highlighted regions that are not inside the text bounding box
    hl_regions = filter_regions_outside_bounding_box(hl_regions, text_bounding_box)
    # Loop over each detected word and its corresponding bounding box
    hl_words = []
    for word in result['words']:
        if word['confidence'] > 0.5:  # Filter out weakly confident text (optional)
            overlap_area = 0
            for region in hl_regions:
                overlap_area += intersect_area(word['bounding_box'], region)
            if overlap_area > area_threshold and sanitize_text(word['text']):
                word['text'] = sanitize_text(word['text'])
                hl_words.append(word)
                print(f"Detected words: {word['text']} Confidence: {word['confidence']}")

    # merge adjacent words
    merged_words = []
    bounding_boxes = []
    for i in range(len(hl_words)):
        if i == 0:
            merged_words.append(hl_words[i]['text'])
            bounding_boxes.append(hl_words[i]['bounding_box'])
        else:
            if hl_words[i]['idx'] - hl_words[i-1]['idx'] == 1:
                merged_words[-1] += ' ' + hl_words[i]['text']
                bounding_boxes[-1] = merge_bounding_boxes([bounding_boxes[-1], hl_words[i]['bounding_box']])
            else:
                merged_words.append(hl_words[i]['text'])
                bounding_boxes.append(hl_words[i]['bounding_box'])

    # for box in hl_regions:
    #     x, y, w, h = box
    #     cv2.rectangle(debug_image, (x, y), (x + w, y + h), (0, 255, 0), 2)

    for box in bounding_boxes:
        x, y, w, h = box
        cv2.rectangle(debug_image, (x, y), (x + w, y + h), (0, 0, 255), 2)

    debug_image_path = processed_image.replace('.jpg', '_debug.jpg')
    cv2.imwrite(debug_image_path, debug_image)
    return {'full_text': result['full_text'], 'highlighted_words': merged_words, 'debug_image': debug_image_path}

def vocabulary_to_doc(vocabulary, output_file, title="Vocabulary"):
    output_doc = Document("template.docx")
    title_paragraph = output_doc.paragraphs[0]
    title_paragraph.text = title  # Set the text of the paragraph directly
    title_paragraph.runs[0].font.size = Pt(16)  # Set font size to 16 points
    add_vocabulary_table(output_doc, vocabulary, title)
    output_doc.save(output_file)

# ==== OCR with LLM ====
def extract_text_from_image(image_path_or_file):
#     prompt = """
# # 任务描述
# 输出图片中所有印刷体的英文文本

# # 定义
# 高亮文本：
# - 背景色为橙色或是黄色的单词短语
# - 但不包含有下划线的文本

# # 输出要求
# - 输出文本的换行与图中保持一致。
# - 对于高亮文本进行加粗（左右两边各添加两个星号)。
# - 如果高亮文��中包含逗号或是句号，则将标点符号两边的单词短语分别进行加粗。
# - 对于其他文本，按原文输出，不添加任何标记。
# - 不要输出除此以外的任何其他内容，如解释、注释等
# - 在最终输出前，再次与原图进行比对，确保高亮文本的加粗符合要求，否则需要进行调整。例如：下划线的文本不应该加粗。
# """
    prompt = """
# Task Description
Output all printed English text in the picture.

# Definition
Highlighted text:
- Words or phrases with an orange or yellow background.
- Underlined or bold text in the picture is not considered as Highlighted text.

# Output Requirements
## Content
- The line breaks in the output text should be the same as in the picture.
- Make sure not missing any text from the picture, especially paragraphs containing highlighted text.
- If the text layout in the picture is 2-column, output text in the left column first, and then output text in the right column.

## Formatting
- For highlighted text, make it bold (adding two asterisks on each side).
- If the highlighted text contains a comma or period, separately bold the words or phrases around the punctuation.
- For other text, output it as it is without any formatting.

## Others
- Do not output any other content such as explanations, comments, etc.
- Before the final output, compare it with the original image again to ensure that the boldness of the highlighted text meets the requirements; 
  otherwise, adjustments need to be made. For example, text with underline should not be bold.
"""
    processed_image = preprocess_image(image_path_or_file)
    result = llm_image_completion(processed_image, prompt)

    if len(result) < 100: # retry if the result is not valid
        result = llm_image_completion(processed_image, f"The request is for student learning and education purpose.\n {prompt}")
        if len(result) < 100: # retry failed
            image_name = image_path_or_file if isinstance(image_path_or_file, str) else image_path_or_file.name
            result = f"Unable to process image {image_path_or_file}"
    return strip_format_quote(result)

def add_markdown_to_docx(doc, markdown_text):
    # Regex to find underlined text (in markdown, __text__)
    lines = markdown_text.split('\n')
    for line in lines:
        asterisked_pattern = re.compile(r'\*\*(.*?)\*\*')
        # Split the markdown text by underlined sections
        parts = asterisked_pattern.split(line)
        
        para = doc.add_paragraph()
        for i, part in enumerate(parts):
            run = para.add_run(part)
            if i % 2 == 1:
                run.font.highlight_color = 7  # 7 represents yellow highlight in python-docx

def images_to_doc(images, output_file, process_callback=None):
    doc = Document()
    for i, img in enumerate(images):
        if process_callback:
            process_callback(i+1, len(images))
        md_text = extract_text_from_image(img)
        add_markdown_to_docx(doc, md_text)
        print(f"--------------------------\n{md_text}")
        doc.add_page_break()
    doc.save(output_file)

def doc_to_markdown(doc):
    paragraphs = []
    for para in doc.paragraphs:
        text = []
        for run in para.runs:
            if run.font.highlight_color:
                text.append(f"**{run.text}**")
            else:
                text.append(run.text)
        paragraphs.append(' '.join(text).strip())
    return '\n\n'.join(paragraphs)

def doc_to_word_group(doc):
    word_groups = []
    for para in doc.paragraphs:
        words = [run.text for run in para.runs if run.font.highlight_color]
        if len(words) > 0:
            word_groups.append({"words": words, "paragraph": para.text})
    return word_groups

def create_vocabulary_from_word_groups(word_groups):
    prompt = dedent(f"""
    对于所有 VOCABULARY 里的单词与短语，在对应的 PARAGRAPH 的上下文语境中翻译。每个单词输出一行，格式为
    单词 | 音标 | 词性 | 中文翻译

    - 如果加粗的单词是动词，则将其转化为动词原形后再翻译输出。例如，加粗的单词为动词lingered，则转成linger。
    - 如果加粗的单词是名词，则将其转化为名词单数形式后再翻译输出。例如，加粗的单词为名词tables，则转成table。
    - 如果是一个词组，则音标与词性可以为空，但 | 不能省略。
    - 如果一个单词在之前已出现过，则跳过该单词不再输出。
    - 单词默认采用小写；专有名词首字母大写。
    - 单词前不需要加序号。
    - 除了上述要求的输出格式以外，不要输出任何其他内容。
    - 在最终输出前，再次检查每个输出的单词或短语，确保符合上述要求，否则进行调整。例如检查是存在动词过去式、名词复数、单词音标或是词性缺失等不符合要求的情况。

    # EXAMPLE
    正确的输出：
    lurch | /lɜːrtʃ/ | v. | 突然倾斜

    错误的输出：
    stranded | /ˈstrændɪd/ | v. | 搁浅
    错误原因：动词没有输出为原形，而是输出了动词的过去式。

    """)
    group_prompt = ""
    for idx, group in enumerate(word_groups):
        words = '\n'.join(group['words'])
        group_prompt += dedent(f"""
    --------------------
    # VOCABULARY {idx+1}
    {words}

    # PARAGRAPH {idx+1}
    {group['paragraph']}

    """)
    result = llm_completion(prompt + group_prompt)
    result = strip_format_quote(result)
    result = [l.strip() for l in result.split('\n') if l.strip()]
    result = [[j.strip() for j in i.split('|')] for i in result]
    return result

def create_vocabulary(words, article):
    words = '\n'.join(words)
    prompt = f"""
    对于 WORDS 中每一行的单词或短语，在 ARTICLE 的上下文语境中翻译。每个翻译输出一行，格式为
    单词 | 音标 | 词性 | 中文翻译

    - 如果单词是动词，则将其转化为动词原形后再翻译输出。例如，加粗的单词为动词lingered，则转成linger。
    - 如果单词是名词，则将其转化为名词单数形式后再翻译输出。例如，加粗的单词为名词tables，则转成table。
    - 如果是一个短语，则音标与词性可以为空，但 | 不能省略。
    - 如果一个单词在之前已出现过，则跳过该单词不再输出。
    - 单词默认采用小写；专有名词首字母大写。
    - 单词前不需要加序号。
    - 除了上述要求的输出格式以外，不要输出任何其他内容。
    - 在最终输出前，再次检查每个输出的单词或短语，确保符合上述要求，否则进行调整。例如检查是存在动词过去式、名词复数、单词音标或是词性缺失等不符合要求的情况。

    # EXAMPLE
    正确的输出：
    lurch | /lɜːrtʃ/ | v. | 突然倾斜

    错误的输出：
    stranded | /ˈstrændɪd/ | v. | 搁浅
    错误原因：动词没有输出为原形，而是输出了动词的过去式。

    # WORDS
    {words}

    # ARTICLE 
    ```
    {article}
    ```
    """
    result = llm_completion(prompt).strip("```")
    result = [l.strip() for l in result.split('\n') if l.strip()]
    result = [[j.strip() for j in i.split('|')] for i in result]
    return result

def add_vocabulary_table(doc, vocabulary, title):
    table = doc.tables[0]

    for word in vocabulary:
        row = table.add_row()
        row.height = Pt(20)
        for j, cell in enumerate(row.cells):
            cell.text = word[j] if j < len(word) else ""

def doc_to_vocabulary(input_file, output_file, title="Vocabulary"):
    doc = Document(input_file)
    word_groups = doc_to_word_group(doc)
    vocabulary = create_vocabulary_from_word_groups(word_groups)

    output_doc = Document("template.docx")
    title_paragraph = output_doc.paragraphs[0]
    title_paragraph.text = title  # Set the text of the paragraph directly
    title_paragraph.runs[0].font.size = Pt(16)  # Set font size to 16 points
    add_vocabulary_table(output_doc, vocabulary, title)
    output_doc.save(output_file)
    return vocabulary

### TESTS LLM ###
def test_convert_markdown_to_docx():
    test_markdown = """
Brenda Z. Guiberson wanted to be a jungle explorer when she was a child. Much of her childhood was spent swimming, watching birds and __salmon__, and searching for arrowheads near her home along the Columbia River in the state of Washington. After volunteering at her child’s school, Guiberson became interested in writing nature books for children. She says that she writes for the child in herself, the one who loves adventure, surprises, and learning new things—a jungle explorer in words.

SETTING A PURPOSE As you read, pay attention to how earthquakes affect people, animals, the land, and the ocean, and think about how people explain and deal with the impact of these damaging events.

Head for the Hills! It’s Earth Against Earth

-----------------------------------------

For centuries, a big __chunk__ of earth under the Indian Ocean known as the India plate has been __scraping__ against another chunk of earth, the Burma plate. At eight o’clock in the morning on December 26, 2004, this scraping reached a breaking point near the island of Sumatra in Indonesia. A 750-mile section of earth snapped and __popped__ up as a new __40-foot-high__ cliff. This created one of the biggest earthquakes ever, 9.2 to 9.3 on the Richter scale.1 At a hospital, oxygen tanks __tumbled__ and beds __lurched__. At a __mosque__, the dome crashed to the floor. On the street, athletes running a race fell

1 Richter scale (rĭk’tar): a scale ranging from 1 to 10 that expresses the amount of energy released by an earthquake; named after Charles Richter, an American seismologist.
"""
    doc = Document()
    add_markdown_to_docx(doc, test_markdown)
    doc.save("test_markdown.docx")

def test_extract_text_from_image():
    test_image_path = "test.jpg"
    print(extract_text_from_image(test_image_path))

def test_images_to_doc():
    images_to_doc(["/Users/ryan/Downloads/IMG_5292.jpg"], "test_markdown.docx")

def test_doc_to_vocabulary():
    # doc = Document("tests/test_doc1.docx")
    # groups = doc_to_word_group(doc)
    # vocabulary = create_vocabulary_from_word_groups(groups)
    # pprint(vocabulary)
    doc_to_vocabulary("tests/test_doc1.docx", "test_vocabulary.docx", "Test Vocabulary")


### TESTS Textract ###
def test_find_highlighted_regions():
    i = 6
    imgutil.preprocess_image(f"tests/test{i}.jpg", f"tests/test{i}_processed.jpg")
    find_highlighted_regions(f"tests/test{i}_processed.jpg", output_path=f"tests/test{i}_highlighted.jpg")

def test_extract_highlighted_words_from_image():
    result = extract_highlighted_words_from_image("tests/walter2.jpg")
    print(result['highlighted_words'])
    debug_image = cv2.imread(result['debug_image'])
    cv2.imshow('Debug Image', debug_image)
    cv2.waitKey(0)
    cv2.destroyAllWindows()

def test_hightlighted_words_to_vocabulary():
    result = extract_highlighted_words_from_image("tests/test1.jpg")
    print("=== Highlighted Words ===")
    print(result['highlighted_words'])
    print("=== Full Text ===")
    print(result['full_text'])
    vocabulary = create_vocabulary(result['highlighted_words'], result['full_text'])
    vocabulary_to_doc(vocabulary, "test_vocabulary.docx")

if __name__ == "__main__":
    # test_extract_text_from_image()
    # test_images_to_doc()
    # test_convert_markdown_to_docx()
    # test_extract_highlighted_words_from_image()
    # test_hightlighted_words_to_vocabulary()
    test_doc_to_vocabulary()
